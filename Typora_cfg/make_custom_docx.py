"""
生成 pandoc 中文文档模板 custom-reference.docx
用法：
  1. pip install python-docx
  2. python make_custom_docx.py
  3. pandoc 你的文件.md -o output.docx --reference-doc="custom-reference.docx"
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 字体常量 ──────────────────────────────────────────────
CN_BODY  = '宋体'       # 正文中文
CN_HEAD  = '黑体'       # 标题中文
EN_FONT  = 'Times New Roman'  # 英文/数字
SZ_BODY  = 11           # 正文字号（pt）
SZ_H1    = 22           # 一级标题
SZ_H2    = 16
SZ_H3    = 14
SZ_CODE  = 9            # 代码块
SZ_TABLE = 9            # 表格


def set_font(run, cn_font=CN_BODY, en_font=EN_FONT, size=SZ_BODY, bold=False, color=None):
    """统一设置 run 的中英文字体、字号、加粗、颜色"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en_font
    # 设置东亚字体（中文）
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(para, before=0, after=0, line=1.15):
    """设置段落间距和行距"""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_style(style, cn_font=CN_BODY, en_font=EN_FONT, size=SZ_BODY,
              bold=False, color=None, align=None, before=0, after=0, line=1.15):
    """修改 Word 内置样式的字体和段落格式"""
    # 段落对齐
    if align is not None:
        style.paragraph_format.alignment = align
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line
    # 字体
    f = style.font
    f.size = Pt(size)
    f.bold = bold
    f.name = en_font
    # 东亚字体
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    if color:
        f.color.rgb = color


def add_table_style(doc):
    """设置表格默认样式：无边框线表格 → 改为有边框"""
    style = doc.styles['Table Grid']
    # 设置表格字体
    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    f = style.font
    f.size = Pt(SZ_TABLE)
    f.name = EN_FONT
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), CN_BODY)


def add_source_code_style(doc):
    """添加或修改代码块样式（仿 GitHub 等宽字体）"""
    try:
        style = doc.styles['Source Code']
    except KeyError:
        style = doc.styles.add_style('Source Code', 1)  # 段落样式
    style.font.name = 'Consolas'
    style.font.size = Pt(SZ_CODE)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.0
    # 淡灰底纹
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F2F2')
    shd.set(qn('w:val'), 'clear')
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style.element.append(pPr)
    pPr.append(shd)


def main():
    doc = Document()

    # ── 1. 修改内置样式 ──────────────────────────────────
    styles = doc.styles

    # 正文 Normal
    set_style(styles['Normal'], CN_BODY, EN_FONT, SZ_BODY,
              before=0, after=4, line=1.35)

    # 各级标题
    set_style(styles['Heading 1'], CN_HEAD, EN_FONT, SZ_H1, bold=True,
              before=18, after=6, line=1.5)
    set_style(styles['Heading 2'], CN_HEAD, EN_FONT, SZ_H2, bold=True,
              before=12, after=4, line=1.3)
    set_style(styles['Heading 3'], CN_HEAD, EN_FONT, SZ_H3, bold=True,
              before=8, after=3, line=1.2)

    # 代码块（用内置的 内联代码 + 自定义 Source Code）
    add_source_code_style(doc)

    # ── 2. 页面设置 ──────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

    # ── 3. 写入示例内容（让模板生效所必需的段落）─────────
    p = doc.add_paragraph('Normal 正文样式示例')
    set_paragraph_spacing(p, before=0, after=6)
    run = p.runs[0]
    set_font(run, CN_BODY, EN_FONT, SZ_BODY)

    h1 = doc.add_heading('Heading 1 一级标题', level=1)
    h2 = doc.add_heading('Heading 2 二级标题', level=2)
    h3 = doc.add_heading('Heading 3 三级标题', level=3)

    # 表格
    table = doc.add_table(rows=3, cols=3, style='Table Grid')
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f'表头{i+1}_{j+1}' if i == 0 else f'数据{i+1}_{j+1}'

    # 代码块占位
    p_code = doc.add_paragraph('code_example()  // 代码块样式占位')
    p_code.style = doc.styles['Source Code']
    set_paragraph_spacing(p_code, before=4, after=4, line=1.0)

    # ── 4. 保存 ──────────────────────────────────────────
    out = 'custom-reference.docx'
    doc.save(out)
    print(f'✅ 已生成模板文件: {out}')
    print(f'   使用: pandoc 你的文件.md -o output.docx --reference-doc="{out}"')


if __name__ == '__main__':
    main()
